import io
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError("Unsupported file format. Please upload PDF, DOCX, or TXT.")
